import os
import tempfile
import logging
import platform
import time
from PIL import Image
from io import BytesIO
from whoosh.fields import Schema, TEXT, ID, STORED, DATETIME
from whoosh.analysis import StemmingAnalyzer, StandardAnalyzer
from whoosh import index
from whoosh.qparser import QueryParser, MultifieldParser, OrGroup
from whoosh.query import Phrase, And, Or, Term
from django.conf import settings
import PyPDF2
from docx import Document

# Configurar el logger
logger = logging.getLogger(__name__)

# pytesseract es opcional - solo para OCR
try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    logger.warning("pytesseract no está disponible. OCR deshabilitado.")

# PyMuPDF (fitz) es opcional
try:
    import fitz
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logger.warning("PyMuPDF no está disponible. Extracción avanzada de PDF deshabilitada.")

# Configuración de Tesseract para Windows
if TESSERACT_AVAILABLE and platform.system() == 'Windows':
    # Ruta específica donde está instalado Tesseract
    tesseract_path = r'C:\Users\jacqu\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'
    tesseract_data_dir = r'C:\Users\jacqu\AppData\Local\Programs\Tesseract-OCR\tessdata'
    
    if os.path.exists(tesseract_path):
        pytesseract.pytesseract.tesseract_cmd = tesseract_path
        os.environ['TESSDATA_PREFIX'] = tesseract_data_dir
        logger.info(f'Tesseract configurado en: {tesseract_path}')
        logger.info(f'Directorio de datos: {tesseract_data_dir}')
    else:
        logger.warning(f'Tesseract no encontrado en {tesseract_path}. La extracción de texto de imágenes no funcionará correctamente.')
elif TESSERACT_AVAILABLE:
    # Para Linux/macOS
    pytesseract.pytesseract.tesseract_cmd = 'tesseract'
    logger.info('Usando Tesseract del PATH del sistema')

# Esquema para el índice de búsqueda
# IMPORTANTE: expediente_id debe ser ID para poder filtrar por él
# Usamos StandardAnalyzer para contenido para mejor búsqueda exacta de párrafos
# y StemmingAnalyzer para título para búsquedas más flexibles
schema = Schema(
    doc_id=ID(stored=True, unique=True),
    expediente_id=ID(stored=True),  # Cambiado a ID para permitir filtrado
    titulo=TEXT(stored=True, analyzer=StemmingAnalyzer()),
    contenido=TEXT(analyzer=StandardAnalyzer(), stored=True),  # Cambiado a StandardAnalyzer y stored=True para mejor búsqueda exacta
    tipo_archivo=STORED,
    fecha_creacion=DATETIME(stored=True),
    ruta_archivo=STORED,
)

def get_index_dir():
    """
    Obtiene el directorio del índice, creándolo si no existe.
    Asegura que el directorio tenga los permisos correctos.
    """
    index_dir = os.path.join(settings.MEDIA_ROOT, 'whoosh_index')
    try:
        if not os.path.exists(index_dir):
            os.makedirs(index_dir, exist_ok=True)
            logger.info(f"Directorio de índices creado en: {index_dir}")
        
        # Verificar permisos de escritura
        test_file = os.path.join(index_dir, '.permission_test')
        with open(test_file, 'w') as f:
            f.write('test')
        os.remove(test_file)
        
        return index_dir
    except Exception as e:
        logger.error(f"Error al acceder al directorio de índices {index_dir}: {str(e)}")
        # Si hay un error, intentar con un directorio temporal
        temp_dir = os.path.join(tempfile.gettempdir(), 'whoosh_index')
        os.makedirs(temp_dir, exist_ok=True)
        logger.warning(f"Usando directorio temporal para índices: {temp_dir}")
        return temp_dir

def get_or_create_index():
    """
    Obtiene o crea un nuevo índice de búsqueda.
    
    Returns:
        whoosh.index.FileIndex: El índice de búsqueda
    """
    try:
        index_dir = get_index_dir()
        
        # Verificar si el directorio existe y es accesible
        if not os.path.exists(index_dir):
            os.makedirs(index_dir, exist_ok=True)
        
        # Verificar si el índice existe y es válido
        if not index.exists_in(index_dir):
            logger.info(f"Creando nuevo índice en: {index_dir}")
            return index.create_in(index_dir, schema)
        
        # Intentar abrir el índice existente
        try:
            return index.open_dir(index_dir)
        except Exception as e:
            logger.warning(f"Error al abrir el índice existente: {str(e)}. Intentando recrear...")
            # Si hay un error al abrir el índice, intentar recrearlo
            import shutil
            backup_dir = f"{index_dir}_backup_{int(time.time())}"
            shutil.move(index_dir, backup_dir)
            logger.info(f"Índice corrupto respaldado en: {backup_dir}")
            return index.create_in(index_dir, schema)
            
    except Exception as e:
        logger.error(f"Error crítico al obtener/crear el índice: {str(e)}")
        # En caso de error crítico, usar un índice en memoria
        logger.warning("Usando índice en memoria como respaldo")
        return index.create_in("mem:", schema)

def extract_text_from_pdf(file_path):
    """Extrae texto de un archivo PDF, con soporte para OCR si es necesario"""
    text = ""
    
    # Método 1: Intentar con PyMuPDF (fitz) - más robusto y rápido
    if PYMUPDF_AVAILABLE:
        try:
            logger.info(f"Extrayendo texto de PDF con PyMuPDF: {file_path}")
            doc = fitz.open(file_path)
            num_pages = len(doc)
            logger.info(f"PDF tiene {num_pages} páginas")
            
            for page_num in range(num_pages):
                page = doc.load_page(page_num)
                page_text = page.get_text() or ""
                text += page_text + "\n"
                if page_text.strip():
                    logger.debug(f"Página {page_num + 1}: {len(page_text)} caracteres extraídos")
            
            doc.close()
            text_stripped = text.strip()
            logger.info(f"Texto extraído con PyMuPDF: {len(text_stripped)} caracteres")
            
            if text_stripped and len(text_stripped) > 10:
                logger.debug(f"Primeros 200 caracteres: {text_stripped[:200]}")
                return text_stripped
            else:
                logger.warning(f"PyMuPDF extrajo poco texto ({len(text_stripped)} caracteres), intentando PyPDF2...")
        except Exception as e:
            logger.warning(f"Error con PyMuPDF: {str(e)}, intentando PyPDF2...")
    
    # Método 2: Intentar con PyPDF2 como respaldo
    try:
        logger.info(f"Extrayendo texto de PDF con PyPDF2: {file_path}")
        text = ""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            num_pages = len(reader.pages)
            logger.info(f"PDF tiene {num_pages} páginas")
            for page_num, page in enumerate(reader.pages, 1):
                try:
                    page_text = page.extract_text() or ""
                    text += page_text + "\n"
                    if page_text.strip():
                        logger.debug(f"Página {page_num}: {len(page_text)} caracteres extraídos")
                except Exception as page_error:
                    logger.warning(f"Error extrayendo página {page_num}: {str(page_error)}")
                    continue
        
        text_stripped = text.strip()
        logger.info(f"Texto extraído con PyPDF2: {len(text_stripped)} caracteres")
        if text_stripped:
            logger.debug(f"Primeros 200 caracteres: {text_stripped[:200]}")
    except Exception as e:
        logger.error(f"Error con PyPDF2: {str(e)}", exc_info=True)
        text_stripped = ""
    
    # Método 3: Si no se pudo extraer suficiente texto, intentar con OCR (si está disponible)
    if (not text_stripped or len(text_stripped) < 50) and TESSERACT_AVAILABLE and PYMUPDF_AVAILABLE:
        logger.info(f"Texto insuficiente ({len(text_stripped)} caracteres), intentando OCR para {file_path}")
        try:
            doc = fitz.open(file_path)
            ocr_text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                # Aumentar la resolución para mejor reconocimiento
                pix = page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72))
                
                # Convertir a imagen PIL
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                
                # Configurar parámetros de Tesseract para mejor reconocimiento
                custom_config = r'--oem 3 --psm 6 -l spa'  # OEM 3 = LSTM + Legacy, PSM 6 = Bloque de texto
                page_text = pytesseract.image_to_string(img, config=custom_config)
                
                if page_text:
                    ocr_text += page_text + "\n"
                    logger.debug(f"OCR página {page_num + 1}: {len(page_text)} caracteres")
            
            doc.close()
            
            if ocr_text.strip():
                logger.info(f"Texto extraído con OCR: {len(ocr_text.strip())} caracteres")
                logger.debug(f"Primeros 200 caracteres OCR: {ocr_text.strip()[:200]}")
                return ocr_text.strip()
        except Exception as ocr_error:
            logger.error(f"Error en OCR: {str(ocr_error)}", exc_info=True)
    
    if not text_stripped:
        logger.warning(f"⚠️ No se pudo extraer texto del PDF {file_path}. Puede ser un PDF escaneado sin OCR disponible.")
    
    return text_stripped

def extract_text_from_docx(file_path):
    """Extrae texto de un archivo DOCX"""
    try:
        logger.info(f"Extrayendo texto de DOCX: {file_path}")
        doc = Document(file_path)
        
        # Extraer texto de párrafos
        paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        
        # También extraer texto de tablas
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = ' '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    table_texts.append(row_text)
        
        # Combinar todo el texto
        all_texts = paragraphs + table_texts
        text = '\n'.join(all_texts)
        
        logger.info(f"Texto extraído de DOCX: {len(text)} caracteres ({len(paragraphs)} párrafos, {len(table_texts)} filas de tabla)")
        if text:
            logger.debug(f"Primeros 200 caracteres: {text[:200]}")
        else:
            logger.warning(f"⚠️ No se encontró texto en el DOCX {file_path}")
        
        return text
    except Exception as e:
        logger.error(f"Error extrayendo texto de DOCX {file_path}: {str(e)}", exc_info=True)
        return ""

def extract_text(file_path):
    """
    Extrae texto de un archivo según su extensión.
    
    Args:
        file_path (str): Ruta al archivo del que extraer texto
        
    Returns:
        str: Texto extraído del archivo, o cadena vacía si no se pudo extraer
    """
    try:
        if not os.path.exists(file_path):
            logger.error(f"El archivo no existe: {file_path}")
            return ""
            
        # Verificar si el archivo está vacío
        if os.path.getsize(file_path) == 0:
            logger.warning(f"El archivo está vacío: {file_path}")
            return ""
            
        ext = os.path.splitext(file_path)[1].lower()
        
        # Intentar determinar el tipo MIME real del archivo
        try:
            import magic
            mime = magic.Magic(mime=True)
            file_type = mime.from_file(file_path)
            logger.debug(f"Tipo MIME detectado para {file_path}: {file_type}")
        except ImportError:
            file_type = ""
            logger.warning("El módulo python-magic no está instalado. No se verificará el tipo MIME real del archivo.")
        
        # Manejar diferentes tipos de archivos
        try:
            if ext == '.pdf':
                return extract_text_from_pdf(file_path)
            elif ext in ['.docx', '.doc']:
                return extract_text_from_docx(file_path)
            elif ext in ['.txt']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            elif ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
                # Para archivos de imagen, usar OCR directamente (si está disponible)
                if TESSERACT_AVAILABLE:
                    try:
                        img = Image.open(file_path)
                        return pytesseract.image_to_string(img, lang='spa')
                    except Exception as img_error:
                        logger.error(f"Error al procesar imagen {file_path} con OCR: {str(img_error)}")
                        return ""
                else:
                    logger.warning(f"OCR no disponible para procesar imagen {file_path}")
                    return ""
            else:
                logger.warning(f"Tipo de archivo no soportado para extracción de texto: {ext}")
                return ""
                
        except Exception as process_error:
            logger.error(f"Error al procesar el archivo {file_path}: {str(process_error)}", exc_info=True)
            return ""
            
    except Exception as e:
        logger.error(f"Error inesperado en extract_text para {file_path}: {str(e)}", exc_info=True)
        return ""

def index_document(documento):
    """Indexa un documento en el motor de búsqueda"""
    try:
        # Validaciones previas antes de indexar
        if not documento.archivo:
            logger.warning(f"Documento {documento.id} no tiene archivo, no se indexará")
            return False
        
        # Verificar que el expediente existe
        if not documento.expediente_id:
            logger.warning(f"Documento {documento.id} no tiene expediente asociado, no se indexará")
            return False
        
        # Obtener la ruta del archivo
        try:
            if hasattr(documento.archivo, 'path'):
                file_path = documento.archivo.path
            else:
                # Si no tiene path, construir la ruta desde el name
                from django.conf import settings
                file_path = os.path.join(settings.MEDIA_ROOT, documento.archivo.name)
        except Exception as path_error:
            logger.error(f"Error obteniendo ruta del archivo para documento {documento.id}: {str(path_error)}")
            return False
        
        # Verificar que el archivo existe físicamente
        if not os.path.exists(file_path):
            logger.error(f"Archivo no encontrado para documento {documento.id}: {file_path}")
            return False
        
        # Verificar que el archivo no está vacío
        if os.path.getsize(file_path) == 0:
            logger.warning(f"Archivo vacío para documento {documento.id}: {file_path}")
            return False
            
        # Extraer texto del archivo
        logger.info(f"Indexando documento {documento.id}: {documento.nombre_documento}")
        logger.info(f"Ruta del archivo: {file_path}")
        logger.info(f"Tamaño del archivo: {os.path.getsize(file_path)} bytes")
        logger.info(f"Tipo de archivo: {documento.tipo_archivo}")
        
        contenido = extract_text(file_path)
        
        # Limpiar y normalizar el texto extraído
        if contenido:
            # Eliminar espacios múltiples y normalizar saltos de línea
            contenido = ' '.join(contenido.split())
            # Asegurar que el contenido tenga al menos algunos caracteres válidos
            contenido = contenido.strip()
        
        # Logging detallado para diagnóstico
        logger.info(f"Texto extraído - Longitud: {len(contenido)} caracteres")
        if contenido and len(contenido) > 0:
            logger.info(f"Texto extraído (primeros 300 caracteres): {contenido[:300]}")
            logger.debug(f"Texto extraído (últimos 100 caracteres): ...{contenido[-100:]}")
        else:
            logger.warning(f"⚠️ No se pudo extraer texto del archivo {file_path}")
            # Si no hay contenido, al menos indexar el nombre del documento para que sea buscable
            contenido = documento.nombre_documento or 'Documento sin título'
            logger.info(f"Usando nombre del documento como contenido indexable: {contenido}")
        
        # Abrir el índice
        ix = get_or_create_index()
        writer = ix.writer()
        
        # Preparar los datos para el índice
        doc_id = f"doc_{documento.id}"
        expediente_id_str = str(documento.expediente_id) if documento.expediente_id else ''
        titulo = documento.nombre_documento or 'Documento sin título'
        contenido_index = contenido or ''
        tipo_archivo = documento.tipo_archivo or ''
        fecha_creacion = documento.fecha_subida
        
        logger.info(f"Guardando en índice - doc_id: {doc_id}, expediente_id: {expediente_id_str}, título: {titulo[:50]}")
        logger.info(f"Contenido a indexar - Longitud: {len(contenido_index)} caracteres")
        
        # Agregar documento al índice
        # IMPORTANTE: expediente_id debe ser string para el índice ID
        writer.update_document(
            doc_id=doc_id,
            expediente_id=expediente_id_str,
            titulo=titulo,
            contenido=contenido_index,
            tipo_archivo=tipo_archivo,
            fecha_creacion=fecha_creacion,
            ruta_archivo=file_path,
        )
        
        writer.commit()
        logger.info(f"✓ Documento {documento.id} indexado correctamente con {len(contenido_index)} caracteres de contenido")
        
        # Verificar que se guardó correctamente
        try:
            with ix.searcher() as searcher:
                stored_doc = searcher.document(doc_id=doc_id)
                if stored_doc:
                    contenido_almacenado = stored_doc.get('contenido', '')
                    logger.info(f"✓ Verificado: Documento guardado en índice con {len(contenido_almacenado)} caracteres")
                else:
                    logger.error(f"✗ Error: Documento {doc_id} no se encontró en el índice después de guardar")
        except Exception as verify_error:
            logger.warning(f"Error al verificar documento en índice: {str(verify_error)}")
        
        return True
    except Exception as e:
        logger.error(f"Error indexando documento {documento.id}: {str(e)}", exc_info=True)
        return False

def remove_document(document_id):
    """Elimina un documento del índice"""
    try:
        ix = get_or_create_index()
        writer = ix.writer()
        writer.delete_by_term('doc_id', f"doc_{document_id}")
        writer.commit()
        return True
    except Exception as e:
        logger.error(f"Error eliminando documento {document_id} del índice: {str(e)}")
        return False

def search_documents(query, expediente_id=None, limit=20, page=1):
    """
    Busca documentos que coincidan con la consulta en el índice de búsqueda.
    Busca en TODOS los documentos del sistema, a menos que se especifique un expediente_id.
    Prioriza coincidencias exactas de palabras, frases y párrafos completos.
    
    Estrategias de búsqueda (enfoque en precisión y flexibilidad):
    - Palabra única: Búsqueda exacta de la palabra, con variaciones solo para palabras largas (>4 caracteres)
    - Frase exacta (con comillas): Busca la frase exacta con flexibilidad mínima para espacios
    - Frase/Párrafo (sin comillas): Busca con múltiples estrategias:
        1. Frase exacta (palabras en orden, sin espacios adicionales) - PRIORIDAD MÁXIMA
        2. Frase con proximidad (permite palabras intermedias según longitud del párrafo)
        3. Para párrafos largos (>5 palabras): también busca que todas las palabras estén presentes
    
    Los resultados se ordenan por relevancia (score), mostrando primero las coincidencias exactas.
    
    IMPORTANTE: Si cambias el esquema del índice, debes reindexar todos los documentos ejecutando:
        python manage.py indexar_documentos --reindexar
    
    Ejemplos:
        - "secretario general" -> Busca la frase exacta "secretario general"
        - secretario general -> Busca primero la frase exacta, luego con proximidad
        - párrafo largo de varias palabras -> Busca con flexibilidad para encontrar el párrafo completo
        - secretario -> Búsqueda exacta de la palabra "secretario"
    
    Args:
        query (str): Términos de búsqueda (puede ser palabra, frase o párrafo)
        expediente_id (int, optional): ID del expediente para filtrar la búsqueda. Si es None, busca en TODOS los documentos.
        limit (int, optional): Número máximo de resultados por página. Por defecto 20.
        page (int, optional): Número de página. Por defecto 1.
        
    Returns:
        dict: Diccionario con los resultados de la búsqueda y metadatos de paginación
    """
    try:
        if not query or not query.strip():
            return {
                'results': [],
                'total': 0,
                'page': page,
                'has_next': False,
                'has_previous': False,
                'page_count': 0,
                'suggestions': []
            }
        
        ix = get_or_create_index()
        
        with ix.searcher() as searcher:
            # Preparar la consulta
            search_query = query.strip()
            
            # Detectar si es una frase exacta (entre comillas) o párrafo/frase (múltiples palabras)
            is_exact_phrase = False
            is_multi_word = len(search_query.split()) > 1
            
            if search_query.startswith('"') and search_query.endswith('"'):
                # Búsqueda de frase exacta entre comillas
                is_exact_phrase = True
                search_query = search_query.strip('"')
            
            # Crear un analizador de consultas que busque en título y contenido
            parser = MultifieldParser(
                ["titulo", "contenido"], 
                ix.schema, 
                group=OrGroup
            )
            
            # Preparar el filtro de expediente (se aplicará después de la búsqueda)
            expediente_id_filter = None
            if expediente_id:
                try:
                    # Convertir a string y validar que no esté vacío
                    expediente_id_str = str(expediente_id).strip()
                    if expediente_id_str and expediente_id_str != 'None' and expediente_id_str != '':
                        expediente_id_filter = expediente_id_str
                except (ValueError, TypeError) as e:
                    logger.warning(f"Error procesando filtro de expediente: {str(e)}")
                    pass
            
            # Analizar la consulta con enfoque en coincidencias exactas
            try:
                if is_multi_word:
                    # Para frases y párrafos, priorizar coincidencias exactas
                    words = search_query.split()
                    num_words = len(words)
                    
                    # Calcular slop dinámico basado en la longitud del párrafo
                    # Para párrafos largos, permitir más flexibilidad
                    if num_words <= 3:
                        # Frases cortas: búsqueda muy exacta
                        slop_exact = 0
                        slop_near = 1
                    elif num_words <= 10:
                        # Frases medianas: permitir algunas palabras intermedias
                        slop_exact = 0
                        slop_near = 2
                    else:
                        # Párrafos largos: permitir más flexibilidad para encontrar el párrafo completo
                        slop_exact = 1
                        slop_near = min(5, num_words // 3)  # Permitir hasta 1/3 de palabras intermedias
                    
                    # Si es una frase exacta (con comillas), buscar la frase exacta con flexibilidad mínima
                    if is_exact_phrase:
                        # Búsqueda de frase exacta, pero con algo de flexibilidad para espacios
                        phrase_exact_content = Phrase("contenido", words, slop=slop_exact)
                        phrase_exact_title = Phrase("titulo", words, slop=slop_exact)
                        q = phrase_exact_content | phrase_exact_title
                    else:
                        # Para frases sin comillas, usar estrategia de coincidencia exacta primero
                        # Estrategia 1: Frase exacta (prioridad máxima)
                        phrase_exact_content = Phrase("contenido", words, slop=slop_exact)
                        phrase_exact_title = Phrase("titulo", words, slop=slop_exact)
                        
                        # Estrategia 2: Frase con proximidad (para párrafos largos)
                        phrase_near_content = Phrase("contenido", words, slop=slop_near)
                        phrase_near_title = Phrase("titulo", words, slop=slop_near)
                        
                        # Estrategia 3: Búsqueda de todas las palabras (para párrafos muy largos)
                        # Esto asegura que si todas las palabras están presentes, se encuentre el documento
                        if num_words > 5:
                            all_words_query = And([Term("contenido", word.lower()) for word in words if len(word) > 2])
                            all_words_title = And([Term("titulo", word.lower()) for word in words if len(word) > 2])
                            # Combinar: priorizar exactas, luego con proximidad, luego todas las palabras
                            q = (phrase_exact_content | phrase_exact_title) | (phrase_near_content | phrase_near_title) | (all_words_query | all_words_title)
                        else:
                            # Combinar: priorizar exactas, luego con proximidad limitada
                            q = (phrase_exact_content | phrase_exact_title) | (phrase_near_content | phrase_near_title)
                    
                else:
                    # Para palabras individuales, buscar coincidencia exacta de la palabra
                    # Primero intentar con búsqueda exacta (sin stemming)
                    search_lower = search_query.lower().strip()
                    
                    # Búsqueda exacta de la palabra completa
                    term_content = Term("contenido", search_lower)
                    term_title = Term("titulo", search_lower)
                    
                    # También buscar variaciones comunes (plural, singular básico)
                    # pero solo si la palabra es lo suficientemente larga
                    if len(search_lower) > 4:
                        # Intentar con el parser para variaciones, pero con peso menor
                        parser_query = parser.parse(search_query)
                        # Combinar: priorizar exactas, luego variaciones
                        q = (term_content | term_title) | parser_query
                    else:
                        # Para palabras cortas, solo búsqueda exacta
                        q = term_content | term_title
                    
            except Exception as parse_error:
                logger.warning(f"Error al analizar la consulta: {str(parse_error)}")
                # Si hay un error, intentar búsqueda más simple
                try:
                    # Buscar términos individuales en contenido
                    terms = search_query.split()
                    queries = []
                    for term in terms:
                        if len(term) > 2:  # Ignorar términos muy cortos
                            term_query = QueryParser("contenido", ix.schema).parse(term)
                            queries.append(term_query)
                    
                    if not queries:
                        return {
                            'results': [],
                            'total': 0,
                            'page': page,
                            'has_next': False,
                            'has_previous': False,
                            'page_count': 0,
                            'suggestions': []
                        }
                    
                    # Combinar todas las consultas con OR
                    q = queries[0]
                    for query_part in queries[1:]:
                        q = q | query_part
                except Exception as e:
                    logger.error(f"Error en búsqueda alternativa: {str(e)}")
                    return {
                        'error': 'Error al procesar la consulta de búsqueda',
                        'details': str(e),
                        'results': [],
                        'total': 0,
                        'page': page,
                        'has_next': False,
                        'has_previous': False,
                        'page_count': 0,
                        'suggestions': []
                    }
            
            # Logging de la consulta de búsqueda
            logger.info(f"Buscando: '{query}' (expediente_id: {expediente_id_filter or 'todos'})")
            logger.debug(f"Consulta Whoosh: {q}")
            
            # Realizar la búsqueda con paginación y ordenamiento por relevancia
            # Usar limit=None para obtener todos los resultados y luego ordenarlos
            start = (page - 1) * limit
            try:
                all_results = searcher.search(q, limit=None, terms=True)
                logger.info(f"Búsqueda realizada: {len(all_results)} resultados encontrados")
                
                # Logging de los primeros resultados para diagnóstico
                if len(all_results) > 0:
                    logger.debug(f"Primeros 3 resultados:")
                    for i, hit in enumerate(all_results[:3], 1):
                        logger.debug(f"  {i}. doc_id: {hit.get('doc_id')}, score: {hit.score:.2f}, título: {hit.get('titulo', '')[:50]}")
                else:
                    logger.warning(f"No se encontraron resultados para: '{query}'")
                    # Intentar búsqueda más simple para diagnóstico
                    simple_query = QueryParser("contenido", ix.schema).parse(query)
                    simple_results = searcher.search(simple_query, limit=5)
                    logger.info(f"Búsqueda simple alternativa encontró {len(simple_results)} resultados")
                    
            except Exception as search_error:
                logger.error(f"Error en la búsqueda de Whoosh: {str(search_error)}", exc_info=True)
                return {
                    'error': 'Error al realizar la búsqueda en el índice',
                    'details': str(search_error),
                    'results': [],
                    'total': 0,
                    'page': page,
                    'has_next': False,
                    'has_previous': False,
                    'page_count': 0,
                    'suggestions': []
                }
            
            # Filtrar por expediente_id si se especificó (después de la búsqueda)
            filtered_results = []
            if expediente_id_filter:
                for hit in all_results:
                    try:
                        hit_expediente_id = hit.get('expediente_id', '')
                        if str(hit_expediente_id) == str(expediente_id_filter):
                            filtered_results.append(hit)
                    except Exception:
                        continue
            else:
                filtered_results = list(all_results)
            
            # Ordenar resultados por relevancia (score) descendente
            # Los resultados con coincidencias exactas tendrán mayor score
            try:
                sorted_results = sorted(filtered_results, key=lambda x: x.score, reverse=True)
            except Exception as sort_error:
                logger.error(f"Error ordenando resultados: {str(sort_error)}", exc_info=True)
                # Si hay error al ordenar, usar los resultados sin ordenar
                sorted_results = filtered_results
            
            # Aplicar paginación manual
            total_results = len(sorted_results)
            paginated_results = sorted_results[start:start + limit]
            
            # Crear un objeto similar a search_page para mantener compatibilidad
            class PaginatedResults:
                def __init__(self, results, page, pagelen, total):
                    self.results = results
                    self.pagenum = page
                    self.pagelen = pagelen
                    self.total = total
                    self.pagecount = (total + pagelen - 1) // pagelen if pagelen > 0 else 0
                
                def __iter__(self):
                    return iter(self.results)
                
                def __len__(self):
                    return len(self.results)
            
            results = PaginatedResults(paginated_results, page, limit, total_results)
            
            # Obtener sugerencias de búsqueda si hay pocos resultados
            suggestions = []
            if len(results) < 3:
                corrector = searcher.corrector("contenido")
                for term in query.split():
                    if len(term) > 3:  # Solo sugerir para términos de más de 3 caracteres
                        suggestions.extend(corrector.suggest(term, limit=3))
            
            # Formatear resultados
            formatted_results = []
            for hit in results:
                try:
                    # Obtener fragmento destacado
                    # Ahora que contenido está stored=True, podemos obtenerlo directamente
                    fragment = ""
                    try:
                        # Intentar obtener highlights del contenido
                        fragment = hit.highlights("contenido", top=3) or ""
                    except Exception:
                        pass
                    
                    # Si no hay highlights, intentar obtener el contenido almacenado y buscar el contexto
                    if not fragment:
                        try:
                            contenido_stored = hit.get('contenido', '')
                            if contenido_stored:
                                # Buscar la primera ocurrencia de alguna palabra de la búsqueda
                                search_words = search_query.lower().split()
                                contenido_lower = contenido_stored.lower()
                                
                                # Encontrar la posición de la primera palabra encontrada
                                for word in search_words:
                                    if len(word) > 2 and word in contenido_lower:
                                        pos = contenido_lower.find(word)
                                        # Extraer contexto alrededor (200 caracteres antes y después)
                                        start = max(0, pos - 200)
                                        end = min(len(contenido_stored), pos + len(word) + 200)
                                        fragment = contenido_stored[start:end]
                                        if start > 0:
                                            fragment = "..." + fragment
                                        if end < len(contenido_stored):
                                            fragment = fragment + "..."
                                        break
                        except Exception:
                            pass
                    
                    # Si aún no hay fragmento, intentar con el título
                    if not fragment:
                        try:
                            fragment = hit.highlights("titulo") or ""
                        except Exception:
                            pass
                    
                    # Si aún no hay fragmento, usar un mensaje genérico
                    if not fragment:
                        fragment = f"Coincidencia encontrada en el documento (relevancia: {hit.score:.2f})"
                    
                    # Asegurarse de que los campos requeridos existan
                    doc_id = hit.get('doc_id', '').replace('doc_', '')
                    titulo = hit.get('titulo', 'Documento sin título')
                    
                    formatted_results.append({
                        'id': doc_id,
                        'titulo': titulo,
                        'fragmento': fragment,
                        'score': hit.score,
                        'fecha_creacion': hit.get('fecha_creacion', ''),
                        'tipo_archivo': hit.get('tipo_archivo', '')
                    })
                except Exception as hit_error:
                    logger.error(f"Error al formatear resultado: {str(hit_error)}")
                    continue
            
            return {
                'results': formatted_results,
                'total': results.total,
                'page': page,
                'has_next': results.pagenum < results.pagecount,
                'has_previous': results.pagenum > 1,
                'page_count': results.pagecount,
                'suggestions': list(set(suggestions))[:5]  # Devolver hasta 5 sugerencias únicas
            }
            
    except Exception as e:
        logger.error(f"Error en la búsqueda: {str(e)}", exc_info=True)
        return {
            'error': 'Error al realizar la búsqueda',
            'details': str(e),
            'results': [],
            'total': 0,
            'page': page,
            'has_next': False,
            'has_previous': False,
            'page_count': 0,
            'suggestions': []
        }
